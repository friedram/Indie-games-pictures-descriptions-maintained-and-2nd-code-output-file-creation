"""
Read.me
What is this? 
Trying to combine code stripping as well as file creation for the 2nd code for reading
in the XboxAppCodeRedeem.
"""

import os
import re
from docx import Document

# Define the regex pattern for Xbox codes (5 groups of 5 uppercase letters/numbers separated by dashes)
CODE_PATTERN = re.compile(r'\b[A-Z0-9]{5}-[A-Z0-9]{5}-[A-Z0-9]{5}-[A-Z0-9]{5}-[A-Z0-9]{5}\b', re.IGNORECASE)

def gather_all_codes(docx_path):
    """
    Gathers *all* codes from:
      - Top-level paragraphs
      - Table cells
    Uses `re.findall` so even if there's extra text in a line (or highlighting),
    we still pick up codes.
    """
    try:
        doc = Document(docx_path)
    except Exception as e:
        print(f"Error opening document '{docx_path}': {e}")
        return []

    found_codes = []

    # 1) Extract from top-level paragraphs
    for para in doc.paragraphs:
        text = para.text.strip()
        codes_in_para = CODE_PATTERN.findall(text)
        found_codes.extend(codes_in_para)

    # 2) Extract from table cells
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    text = para.text.strip()
                    codes_in_para = CODE_PATTERN.findall(text)
                    found_codes.extend(codes_in_para)

    return found_codes

def pick_every_second_code(codes):
    """
    From the list of all matched codes, keep only the 2nd, 4th, 6th, etc.
    """
    return [code for i, code in enumerate(codes) if i % 2 == 1]

def write_codes_to_new_docx(codes, source_path, suffix):
    """
    Creates a new .docx file with the specified suffix appended to the base filename
    and writes each code on its own paragraph.
    """
    base, ext = os.path.splitext(source_path)
    output_path = f"{base}_{suffix}{ext}"  # Example: "D:\path\to\file_2ndCode.docx"

    new_doc = Document()
    for code in codes:
        new_doc.add_paragraph(code)
    
    new_doc.save(output_path)
    return output_path

def remove_pattern_from_docx(input_path, output_path):
    """
    Removes all Xbox codes from the document while **preserving images, tables, and formatting**.
    Uses paragraph `runs` to modify text without losing images.
    """
    try:
        doc = Document(input_path)
    except Exception as e:
        print(f"Error opening document '{input_path}': {e}")
        return

    # Remove only the codes while keeping images and document structure
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if re.search(CODE_PATTERN, run.text):
                run.text = re.sub(CODE_PATTERN, '', run.text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        if re.search(CODE_PATTERN, run.text):
                            run.text = re.sub(CODE_PATTERN, '', run.text)

    # Save the modified document without affecting images
    doc.save(output_path)

def main():
    input_file = input("Enter the full path to the .docx file: ").strip()
    
    if not os.path.exists(input_file):
        print(f"Error: The file '{input_file}' does not exist.")
        return
    
    # Extract all codes
    all_codes = gather_all_codes(input_file)
    if not all_codes:
        print("No codes found in the document.")
        return
    
    print("\nAll codes found (in the order encountered):")
    for c in all_codes:
        print(" -", c)
    
    # Keep every second code
    chosen_codes = pick_every_second_code(all_codes)
    if not chosen_codes:
        print("\nNo codes remain after picking 2nd, 4th, 6th, etc.")
        return
    
    print("\nCodes chosen (2nd, 4th, 6th...):")
    for c in chosen_codes:
        print(" -", c)

    # Create new documents with the stripped text and every second code
    stripped_doc = f"{os.path.splitext(input_file)[0]}_Stripped.docx"
    second_code_doc = f"{os.path.splitext(input_file)[0]}_2ndCode.docx"

    # Write only the selected second codes
    write_codes_to_new_docx(chosen_codes, input_file, "2ndCode")

    # Remove all codes from the original document and save as "Stripped"
    remove_pattern_from_docx(input_file, stripped_doc)

    print(f"\nStripped document saved as: {stripped_doc}")
    print(f"2nd code document saved as: {second_code_doc}")

if __name__ == "__main__":
    main()
